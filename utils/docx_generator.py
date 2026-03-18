from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from datetime import datetime


def generate_import_docx(
    source_name: str,
    url: str,
    worker_results: list,
    clean_data: dict,
    output_path: Path,
) -> None:
    """
    Generate a Data Commons Import Document (.docx).
    Uses LLM-compiled clean_data if available, falls back to raw worker results.
    """
    doc = Document()

    # Helper to get worker result by name
    def get_raw(name: str) -> str:
        for wr in worker_results:
            if wr["name"] == name:
                return wr.get("result", "").strip()
        return ""

    # Helper to get clean value or fall back to raw
    def get(key: str, fallback_worker: str = "") -> str:
        if clean_data and clean_data.get(key) and clean_data[key] != "Not found":
            return str(clean_data[key]).strip()
        if fallback_worker:
            raw = get_raw(fallback_worker)
            return raw if raw else "Not found"
        return "Not found"

    # ─── Title ───
    title = doc.add_paragraph("Data Commons Import", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(source_name, style="Title")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Author + Date
    doc.add_paragraph(f"Author: [Auto-generated]")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%B %Y')}")

    # ─── About ───
    about_p = doc.add_paragraph()
    about_p.add_run("About").bold = True
    doc.add_paragraph(get("about", "A2"))

    # ─── Data Source Summary ───
    summary_p = doc.add_paragraph()
    summary_p.add_run("Data Source Summary").bold = True
    doc.add_paragraph(get("data_source_summary", "B7"))

    # ─── Main Table ───
    table_rows = [
        ("Link to dataset preview or raw data", get("link_to_data", "B6") or url),
        ("Link that best explains the dataset", get("link_explains_dataset", "B10.1") or url),
        ("Place types covered", get("place_types", "A3")),
        ("Place ID resolution", get("place_id_resolution", "C14.2")),
        ("Date range covered", get("date_range")),
        ("Number of variables", get("num_variables", "A1")),
        (
            "Are any top-level variables (e.g., w/ single constraint) missing?\n\n"
            "If yes, can they be aggregated?",
            get("top_level_missing"),
        ),
        ("Refresh Cycle", get("refresh_cycle", "C15")),
        ("Code/data location and import name", get("code_data_location", "B6.1")),
        ("License", get("license", "A4")),
        ("Approver", ""),
    ]

    # Build date_range from raw workers if not in clean_data
    if table_rows[4][1] == "Not found":
        min_d = get_raw("C12")
        max_d = get_raw("C13")
        date_range = _build_date_range(min_d, max_d)
        table_rows[4] = ("Date range covered", date_range)

    table = doc.add_table(rows=len(table_rows), cols=2, style="Table Grid")
    for i, (key, value) in enumerate(table_rows):
        key_cell = table.rows[i].cells[0]
        val_cell = table.rows[i].cells[1]
        key_cell.text = key
        val_cell.text = value
        for run in key_cell.paragraphs[0].runs:
            run.bold = True

    doc.add_paragraph()

    # ─── Statistical Variable Categories ───
    doc.add_heading("Statistical Variable Categories", level=1)
    doc.add_paragraph(
        "List the categories of variables. Include distinct populationType values, "
        "measuredProperty values and constraint properties. For each constraint "
        "property, list only the value enum type, not instances."
    )
    stat_vars = get("stat_var_categories", "A1")
    if stat_vars and stat_vars != "Not found":
        doc.add_paragraph(stat_vars)

    # ─── Examples ───
    doc.add_heading("Examples (not necessary for first pass review)", level=1)
    doc.add_paragraph(
        "Provide examples of no more than 3 StatVars. Please ensure the DCIDs are "
        "constructed using the suggested naming conventions."
    )

    # ─── Notes and Caveats ───
    doc.add_heading("Notes and Caveats (optional)", level=1)
    notes = get("notes")
    if notes and notes != "Not found":
        doc.add_paragraph(notes)
    else:
        doc.add_paragraph("No additional notes.")

    doc.save(str(output_path))


def _build_date_range(min_date: str, max_date: str) -> str:
    min_d = min_date if min_date and "not found" not in min_date.lower() else ""
    max_d = max_date if max_date and "not found" not in max_date.lower() else ""
    if min_d and max_d:
        return f"{min_d} – {max_d}"
    return min_d or max_d or "Not found"
